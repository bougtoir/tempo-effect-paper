# -*- coding: utf-8 -*-
"""Create cover letters for Lancet Correspondence and PDR Research Note (EN + JP)."""
import os

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def make_cover_letter(path, paragraphs, font_name='Times New Roman', font_size=12, line_spacing=1.5):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(font_size)
    style.paragraph_format.line_spacing = line_spacing

    for text, kwargs in paragraphs:
        p = doc.add_paragraph()
        bold = kwargs.get('bold', False)
        italic = kwargs.get('italic', False)
        size = kwargs.get('size', font_size)
        align = kwargs.get('align', None)
        space_after = kwargs.get('space_after', 6)
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if align:
            p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
    doc.save(path)
    print(f"OK: {path}")


SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = os.path.join(os.path.dirname(SCRIPT_DIR), 'manuscripts')
os.makedirs(OUT_DIR, exist_ok=True)

today = "April 14, 2026"
R = WD_ALIGN_PARAGRAPH.RIGHT

# =====================================================================
# 1. Lancet Correspondence -- English
# =====================================================================
lancet_en = [
    (today, {'align': R, 'space_after': 18}),
    ("The Editor\nThe Lancet\n125 London Wall\nLondon EC2Y 5AS, United Kingdom",
     {'space_after': 18}),
    ("Dear Editor,", {'space_after': 12}),
    ('Re: Submission of Correspondence \u2014 "The Forgotten Tempo Effect: How Delayed '
     'Childbearing Shapes Simultaneously Living Population Size Across OECD Countries"',
     {'bold': True, 'space_after': 12}),
    ("We wish to submit the enclosed Correspondence for consideration in The Lancet. "
     "This letter presents a concise, policy-relevant finding: the timing of births (tempo effect) "
     "exerts an independent and quantitatively large influence on the number of people simultaneously "
     "alive, yet this mechanism is almost entirely absent from current pronatalist policy discourse.",
     {'space_after': 8}),
    ("Using a parsimonious endogenous renewal model validated against UN World Population Prospects "
     "2024 data for 40 countries (38 OECD members plus China and the DRC) over 1970\u20132023, we show "
     "that a 5-year increase in mean age at childbearing reduces simultaneously living population by "
     "approximately one-sixth, independent of total fertility rate. The dynamic model achieves a "
     "median absolute percentage error of 4.6% against observed trajectories \u2014 demonstrating that "
     "this simple demographic identity captures population dynamics with surprising accuracy.",
     {'space_after': 8}),
    ("The Lancet\u2019s readership of clinicians, public health specialists, and policymakers is ideally "
     "positioned to appreciate the implications of this finding. While Bongaarts and Feeney (1998) "
     "and Goldstein, Lutz, and Scherbov (2003) established the theoretical foundations, the tempo "
     "effect has since been largely forgotten in policy circles. Our 40-country empirical validation "
     "brings renewed, evidence-based urgency to this overlooked lever.",
     {'space_after': 8}),
    ("The manuscript is approximately 300 words with one figure, conforming to The Lancet\u2019s "
     "Correspondence format requirements. The work is original, has not been published previously, "
     "and is not under consideration elsewhere. All authors have approved the manuscript and "
     "declare no competing interests.",
     {'space_after': 8}),
    ("We believe this Correspondence will be of significant interest to The Lancet\u2019s global "
     "readership given the widespread policy concern about demographic decline across OECD countries.",
     {'space_after': 12}),
    ("Yours sincerely,", {'space_after': 18}),
    ("Tatsuki Onishi\n[Affiliation]\n[Email]\n[ORCID]",
     {'space_after': 6}),
]
make_cover_letter(os.path.join(OUT_DIR, 'CoverLetter_Lancet_EN.docx'), lancet_en)

# =====================================================================
# 2. Lancet Correspondence -- Japanese
# =====================================================================
lancet_jp = [
    (today, {'align': R, 'space_after': 18}),
    ("The Editor\nThe Lancet\n125 London Wall\nLondon EC2Y 5AS, United Kingdom",
     {'space_after': 18}),
    ("Dear Editor,", {'space_after': 12}),
    ('件名: Correspondence投稿 \u2014 "The Forgotten Tempo Effect: How Delayed '
     'Childbearing Shapes Simultaneously Living Population Size Across OECD Countries"',
     {'bold': True, 'space_after': 12}),
    ("同封のCorrespondenceをThe Lancetへの掲載候補としてご検討いただきたく、投稿いたします。"
     "本レターは、政策的に重要な知見を簡潔に提示するものです：出生のタイミング（テンポ効果）が、"
     "同時に生存している人口数に対して独立かつ定量的に大きな影響を及ぼすにもかかわらず、"
     "現行の少子化対策の議論からほぼ完全に欠落しているという事実です。",
     {'space_after': 8}),
    ("国連世界人口推計2024のデータを用いて40カ国（OECD加盟38カ国＋中国・コンゴ民主共和国）について"
     "1970〜2023年の期間で検証した簡素な内生更新モデルにより、平均出産年齢（MAC）の5年上昇が、"
     "合計特殊出生率（TFR）とは独立に、同時在生人口を約6分の1減少させることを示しました。"
     "動的モデルは観測値に対して中央値絶対パーセント誤差4.6%を達成しており、"
     "この単純な人口学的恒等式が驚くほどの精度で人口動態を再現することを実証しています。",
     {'space_after': 8}),
    ("The Lancetの読者層である臨床医、公衆衛生専門家、政策立案者は、"
     "この知見の含意を理解する最適な立場にあります。"
     "Bongaarts & Feeney (1998) およびGoldstein, Lutz, & Scherbov (2003) が理論的基盤を確立して以来、"
     "テンポ効果は政策的議論においてほぼ忘れ去られてきました。"
     "本稿の40カ国実証検証は、この見過ごされた政策レバーにエビデンスに基づく新たな緊急性をもたらします。",
     {'space_after': 8}),
    ("本原稿は約300語・図1点で、The LancetのCorrespondence投稿規定に準拠しています。"
     "本研究はオリジナルであり、他誌に掲載済みまたは投稿中ではありません。"
     "全著者が原稿を承認しており、利益相反はありません。",
     {'space_after': 8}),
    ("OECD諸国における人口減少への政策的関心が広がる中、本CorrespondenceがThe Lancetの"
     "グローバルな読者にとって重要な関心事となると確信しております。",
     {'space_after': 12}),
    ("敬具", {'space_after': 18}),
    ("大西 達己\n[所属機関]\n[メールアドレス]\n[ORCID]",
     {'space_after': 6}),
]
make_cover_letter(os.path.join(OUT_DIR, 'CoverLetter_Lancet_JP.docx'), lancet_jp)

# =====================================================================
# 3. PDR Research Note -- English
# =====================================================================
pdr_en = [
    (today, {'align': R, 'space_after': 18}),
    ("The Editors\nPopulation and Development Review\nPopulation Council\n"
     "One Dag Hammarskjold Plaza\nNew York, NY 10017, USA",
     {'space_after': 18}),
    ("Dear Editors,", {'space_after': 12}),
    ('Re: Submission of Notes and Commentary \u2014 "The Forgotten Tempo Effect: Delayed '
     'Childbearing, Simultaneously Living Population, and the Pace of Social Adaptation '
     'Across OECD Countries"',
     {'bold': True, 'space_after': 12}),
    ("We wish to submit the enclosed manuscript for consideration as a Notes and Commentary "
     "article in Population and Development Review. This paper revisits the tempo effect \u2014 "
     "the independent influence of birth timing on population size \u2014 which, despite seminal "
     "contributions by Bongaarts and Feeney (1998) and the foundational work by Goldstein, "
     "Lutz, and Scherbov (2003) published in this journal, has largely disappeared from "
     "contemporary policy discourse on demographic decline.",
     {'space_after': 8}),
    ("We believe PDR is the natural venue for this work for three reasons:\n\n"
     "1. Intellectual lineage: Goldstein, Lutz, and Scherbov\u2019s (2003) demonstration that "
     "generational length changes reduce population size was published in PDR. Our paper "
     "extends their EU-15 analysis to 40 countries with 20 additional years of data, using "
     "a complementary modelling approach.\n\n"
     "2. Empirical contribution: Using a parsimonious four-parameter endogenous renewal model "
     "validated against UN WPP 2024 data for 38 OECD member states plus China and the DRC "
     "over 1970\u20132023, we achieve a median absolute percentage error of 4.6% \u2014 demonstrating "
     "that quantum, tempo, and survival alone explain the vast majority of population dynamics "
     "across diverse demographic contexts.\n\n"
     "3. Policy reframing: We introduce the concept that tempo-sensitive policies control not "
     "merely population size but the speed at which societies must adapt their institutions "
     "to demographic change. This \u201cpace of adaptation\u201d framing \u2014 showing that a 5-year MAC "
     "increase reduces simultaneously living population by ~1/6 independent of TFR \u2014 offers "
     "a novel perspective for the PDR readership.",
     {'space_after': 8}),
    ("The manuscript is approximately 6,500 words with 7 figures, 4 tables, and 3 appendices "
     "(including a comparative table of national population projection methodologies for 15 "
     "countries/agencies). It follows GATHER reporting guidelines. The manuscript is formatted "
     "for double-anonymised review; author-identifying information appears only in this cover "
     "letter.",
     {'space_after': 8}),
    ("Disclosure: A preliminary summary of these findings (~300 words) was previously submitted to "
     "The Lancet as a Correspondence but was not accepted for publication. That short communication "
     "presented the core result without the full model specification, 40-country validation, bias "
     "analysis, or policy framework developed here. We disclose this prior submission in the interest "
     "of transparency. The present manuscript represents a substantially expanded and independent work.",
     {'space_after': 8}),
    ("The work is original and is not under consideration for publication as a full article "
     "elsewhere. All authors have approved the manuscript and declare no competing interests.",
     {'space_after': 12}),
    ("Yours sincerely,", {'space_after': 18}),
    ("Tatsuki Onishi\n[Affiliation]\n[Email]\n[ORCID]",
     {'space_after': 6}),
]
make_cover_letter(os.path.join(OUT_DIR, 'CoverLetter_PDR_EN.docx'), pdr_en)

# =====================================================================
# Demographic Research -- English Cover Letter
# =====================================================================
demres_en = [
    (today, {'align': R, 'space_after': 18}),
    ("The Editors\nDemographic Research\nMax Planck Institute for Demographic Research\n"
     "Konrad-Zuse-Str. 1\n18057 Rostock, Germany",
     {'space_after': 18}),
    ("Dear Editors,", {'space_after': 12}),
    ('Re: Submission of Research Article \u2014 "The Forgotten Tempo Effect: Delayed '
     'Childbearing, Simultaneously Living Population, and the Pace of Social Adaptation '
     'Across OECD Countries"',
     {'bold': True, 'space_after': 12}),
    ("We wish to submit the enclosed manuscript for consideration as a Research Article "
     "in Demographic Research. This paper revisits the tempo effect \u2014 the independent "
     "influence of birth timing on population size \u2014 which, despite seminal contributions "
     "by Bongaarts and Feeney (1998) and foundational work by Goldstein, Lutz, and "
     "Scherbov (2003), has largely disappeared from contemporary policy discourse on "
     "demographic decline.",
     {'space_after': 8}),
    ("We believe Demographic Research is the ideal venue for this work for three reasons:\n\n"
     "1. Intellectual home: Demographic Research has been central to the development of "
     "tempo-adjustment methodology, publishing key contributions including Kohler and "
     "Ortega (2002) on tempo-adjusted parity progression measures and Frejka and Sobotka "
     "(2008) on diverse European fertility patterns. Our paper extends this tradition by "
     "providing the first systematic 40-country empirical validation of the tempo effect on "
     "simultaneously living population.\n\n"
     "2. Empirical contribution: Using a parsimonious four-parameter endogenous renewal model "
     "validated against UN WPP 2024 data for 38 OECD member states plus China and the DRC "
     "over 1970\u20132023, we achieve a median absolute percentage error of 4.6% \u2014 demonstrating "
     "that quantum, tempo, and survival alone explain the vast majority of population dynamics "
     "across diverse demographic contexts.\n\n"
     "3. Policy reframing: We introduce the concept that tempo-sensitive policies control not "
     "merely population size but the speed at which societies must adapt their institutions "
     "to demographic change. This \u2018pace of adaptation\u2019 framing \u2014 showing that a 5-year MAC "
     "increase reduces simultaneously living population by approximately one-sixth, independent "
     "of TFR \u2014 offers a novel perspective that aligns with the journal\u2019s commitment to "
     "publishing research with clear policy relevance.",
     {'space_after': 8}),
    ("The manuscript is approximately 6,500 words with 7 figures, 4 tables, and 3 appendices "
     "(including a comparative table of national population projection methodologies for 15 "
     "countries/agencies, and a natural experiments analysis examining countries with major "
     "political/border changes as exogenous shocks). It follows GATHER reporting guidelines "
     "and the Demographic Research author guidelines. All analytical code and data sources "
     "are documented for full reproducibility. The manuscript is formatted for double-anonymised "
     "review; author-identifying information appears only in this cover letter.",
     {'space_after': 8}),
    ("開示事項：本研究の予備的要約（約300語）をThe LancetにCorrespondenceとして投稿しましたが、"
     "掲載には至りませんでした。当該短報は、本稿で展開する完全なモデル仕様、40カ国検証、バイアス分析、"
     "政策フレームワークを含みません。透明性の観点からこの投稿歴を開示するものです。"
     "本原稿は大幅に拡張された独立した研究です。",
     {'space_after': 8}),
    ("本研究はオリジナルであり、フル論文として他誌に投稿中ではありません。"
     "全著者が原稿を承認しており、利益相反はありません。",
     {'space_after': 12}),
    ("Yours sincerely,", {'space_after': 18}),
    ("Tatsuki Onishi\n[Affiliation]\n[Email]\n[ORCID]",
     {'space_after': 6}),
]
make_cover_letter(os.path.join(OUT_DIR, 'CoverLetter_PDR_JP.docx'), pdr_jp)

print("Cover letters created.")
