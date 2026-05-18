#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Create PDR Research Note -- Japanese version.

PDR Author Guidelines compliance:
  - Notes and Commentary (8,000-10,000 words)
  - Double-anonymised review
  - Author-date in-text citations
  - Alphabetical reference list
  - 12 pt Times New Roman, double-spaced, 1-inch margins
  - Figures/tables inline after first-mention paragraph
  - GATHER reporting compliance
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
FIG_DIR = os.path.join(os.path.dirname(SCRIPT_DIR), 'figures')
OUT_DIR = os.path.join(os.path.dirname(SCRIPT_DIR), 'manuscripts')
os.makedirs(OUT_DIR, exist_ok=True)


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def add_para(doc, text, bold=False, italic=False, size=12, align=None,
             space_after=6, space_before=0, first_line_indent=None,
             font_name=None):
    """Add a paragraph with a single run."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if font_name:
        run.font.name = font_name
    if align is not None:
        p.alignment = align
    fmt = p.paragraph_format
    fmt.space_after = Pt(space_after)
    fmt.space_before = Pt(space_before)
    if first_line_indent is not None:
        fmt.first_line_indent = Pt(first_line_indent)
    return p


def add_heading_styled(doc, text, level=1):
    """Add a heading with black font colour."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_figure(doc, path, caption, width=6.0):
    """Insert a figure image followed by its caption."""
    if not os.path.exists(path):
        add_para(doc, f'[Figure image not found: {path}]',
                 italic=True, size=10)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    p.paragraph_format.space_before = Pt(12)
    # caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = cap.add_run(caption)
    r.font.size = Pt(10)
    r.italic = True
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(12)


def add_table_caption(doc, caption):
    """Add a table caption above a table."""
    p = doc.add_paragraph()
    r = p.add_run(caption)
    r.font.size = Pt(10)
    r.bold = True
    r.italic = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def set_cell_font(cell, size=10):
    """Set font size for all runs in a table cell."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(size)


# ===========================================================================
# BUILD DOCUMENT
# ===========================================================================
doc = Document()

# Page setup: 1-inch margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# Default style: Times New Roman 12 pt, double-spaced
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0

# ===== TITLE PAGE =====
add_para(doc, "NOTES AND COMMENTARY", bold=True, size=10,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

add_para(
    doc,
    "忘れられたテンポ効果：出産の遅延、同時在生人口、\n"
    "そしてOECD諸国における社会適応速度の制御",
    bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18,
)

add_para(
    doc,
    "[匿名査読のため著者名を削除]",
    italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6,
)

add_para(
    doc,
    "語数：約5,500語（参考文献・付録を除く）",
    italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24,
)

# ===== ABSTRACT =====
add_heading_styled(doc, '要旨', level=2)

add_para(
    doc,
    "人口予測と少子化対策は、出生率のカンタム（出生数）に圧倒的に焦点を当て、"
    "テンポ効果——出産タイミングが同時に生存する世代数に与える独立した"
    "影響——を見落としてきた。本稿では、年齢別出生率（平均出産年齢MACを"
    "中心とする正規分布）とGompertzパラメトリック生存関数（平均寿命に"
    "キャリブレーション）を結合した簡素な内生更新モデルにより、この"
    "過小評価されたメカニズムを再考する。国連世界人口推計2024のデータを"
    "用い、OECD加盟38カ国＋中国＋コンゴ民主共和国の計40カ国で1970–2023年の"
    "検証を行った結果、動的モデル（10年ごとにパラメータ更新）は観測"
    "された人口軌跡に対し絶対百分率誤差の中央値4.6%を達成した。MACの"
    "5年上昇は、TFRとは独立に同時在生人口を約6分の1減少させることを"
    "示す。この知見は、出生数のみに対処する少子化対策がその人口学的"
    "影響を体系的に過大評価すること、そしてテンポに敏感な介入——"
    "住宅、保育、教育改革——が人口転換のペースと社会構造的適応の"
    "速度を管理する未活用のレバーとなりうることを意味する。",
    size=12, space_after=12,
)

add_para(
    doc,
    "キーワード：テンポ効果、同時在生人口、第一子出産年齢、"
    "平均出産年齢、Gompertz生存、OECD、人口転換、人口予測",
    italic=True, size=10, space_after=18,
)

doc.add_page_break()

# ===================================================================
# SECTION 1
# ===================================================================
add_heading_styled(doc, '1. 忘れられたテンポ効果', level=1)

add_para(
    doc,
    "人口減少をめぐるグローバルな議論は、単一の指標——合計特殊出生率"
    "（TFR）——に支配されている。TFRが置換水準（女性１人当たり約"
    "2.1人）を下回ると警鐘が鳴る。しかしこのフレーミングは、"
    "ある瞬間に同時に生存する人数を決定する第二の独立した人口学的"
    "力——出産のタイミング——を体系的に見落としている。",
    size=12, space_after=12,
)

add_para(
    doc,
    "BongaartsとFeeney（1998）は、出生率のカンタム（出生数）とテンポ"
    "（出産タイミング）の区別を定式化し、女性が出産を先送りすると"
    "——たとえコーホート完結出生率が変わらなくても——期間TFRが"
    "機械的に押し下げられることを示した。Goldstein, Lutz, Scherbov（2003）は"
    "この知見をさらに進め、EU15カ国において出産の遅延がある時点で"
    "同時に生存する世代数を減少させ、女性１人当たりの出生数とは"
    "独立に人口減少をもたらすことを実証した。彼らの要因分解は、"
    "世代長の変化が予測される人口減少のかなりの部分を説明する"
    "ことを示した。",
    size=12, space_after=12,
)

add_para(
    doc,
    "この基礎的研究にもかかわらず、テンポの次元は現代の政策議論"
    "からほぼ姿を消している。韓国の記録的な47兆ウォンの少子化対策"
    "から日本の歴代『少子化社会対策大綱』に至るまで、OECD諸国の"
    "最近の少子化対策パッケージを概観すると、出生数の増加にほぼ"
    "排他的に焦点を当てていることが分かる。同時在生人口（SLP）"
    "——ある暦上の時点で生存する人の総数——はカンタムとテンポ"
    "の両方により形成されるが、体系的な政策的関心を受けるのは"
    "カンタムだけである。",
    size=12, space_after=12,
)

add_para(
    doc,
    "本稿は、単純だが実証的に根拠のある人口モデルを用いてテンポ"
    "効果を再考し、３つの問いに答える。（1）テンポ・カンタム・"
    "生存の三要素で構築された簡素なモデルは観測された人口軌跡を"
    "再現できるか？（2）多様な人口学的文脈において、テンポ効果は"
    "カンタムに対してどの程度大きいか？（3）社会が人口変動に合わ"
    "せて制度を適応させるスピードに対する含意は何か？",
    size=12, space_after=12,
)

# ===================================================================
# SECTION 2
# ===================================================================
add_heading_styled(doc, '2. モデルとデータ', level=1)
add_heading_styled(doc, '2.1 内生更新モデル', level=2)

add_para(
    doc,
    "離散時間・単一性別の人口モデルを構築する。人口ベクトル"
    "P(t) = [P₀(t), P₁(t), …, P₁₀₀(t)] は毎年以下のように更新される。",
    size=12, space_after=6,
)

add_para(
    doc,
    "(a) 生存：年齢xの個人はGompertzハザード関数 "
    "h(x) = a·exp(b·x) から導かれる生存確率で年齢x+1まで"
    "生存する。生存関数は S(x) = exp[−(a/b)(exp(bx)−1)]。"
    "パラメータaは出生時平均余命"
    "e₀ = ∫₀∞ S(x)dx が観測値と一致するよう"
    "キャリブレーション。bは0.085に固定。",
    size=12, space_after=6,
)

add_para(
    doc,
    "(b) 出生：出生は内生的に生成。年齢別出生率（ASFR）は"
    "平均出産年齢（MAC）を中心とし標準偏差σの正規密度を"
    "TFRにスケーリング。時刻tの出生数は "
    "Σ(x=15〜49) P_x(t) · f · ASFR(x)。",
    size=12, space_after=6,
)

add_para(
    doc,
    "この最小限のパラメータ化には期間あたり4つの入力値のみが"
    "必要：TFR、平均寿命（e₀）、MAC、σ。モデルは移民を意図的に"
    "省略し、カンタム・テンポ・生存の純粋な人口学的メカニズムを"
    "分離する。",
    size=12, space_after=12,
)

add_heading_styled(doc, '2.2 データ', level=2)

add_para(
    doc,
    "全入力パラメータと検証データは国連世界人口推計2024"
    "（United Nations 2024）から取得。分析対象は40カ国："
    "OECD全加盟38カ国＋中国＋コンゴ民主共和国（DRC）。"
    "初期人口年齢構造（5歳階級）は1歳刻みに内挿。"
    "TFR、e₀、MACは1950–2023年の各暦年について抽出。",
    size=12, space_after=6,
)

add_para(
    doc,
    "GATHER報告ガイドライン（Stevens et al. 2016）に従い："
    "入力データは国連人口部から公開入手可能。"
    "全モデルコードとパラメータは文書化済み。"
    "分析アプローチは完全に再現可能。",
    size=12, space_after=12,
)

add_heading_styled(doc, '2.3 モデルのバリアント', level=2)

add_para(doc, "2つのバリアントを実装する。", size=12, space_after=6)

add_para(
    doc,
    "静的モデル：パラメータ（TFR、e₀、MAC）を基準年の値に固定。"
    "4つの基準年（1970、1980、1990、2000）で前方投影し"
    "160の国×基準年の組み合わせを得る。",
    size=12, space_after=6,
)

add_para(
    doc,
    "動的モデル：10年ごとに観測値を用いてパラメータを更新し、"
    "全40カ国で1970年から2023年まで実行。",
    size=12, space_after=12,
)

# ===================================================================
# SECTION 3
# ===================================================================
add_heading_styled(doc, '3. OECD全体の検証結果', level=1)
add_heading_styled(doc, '3.1 全体的な適合度', level=2)

add_para(
    doc,
    "表1にモデル性能を要約する。動的モデルは53年間で"
    "MAPE中央値4.6%（平均6.7%）を達成し、最終人口比の平均は"
    "0.999（SD=0.189）——体系的バイアスが無視できることを示す。"
    "静的モデルの適合度は予測期間に伴い劣化する："
    "中央値4.7%（基準年2000）から7.3%（基準年1970）。",
    size=12, space_after=6,
)

add_para(
    doc,
    "40カ国のうち30カ国で動的MAPEが10%未満、"
    "20カ国で5%未満、6カ国で2%未満。"
    "最適合国——フランス（0.4%）、コスタリカ（0.9%）、"
    "フィンランド（0.8%）、チェコ（1.3%）、"
    "スロベニア（1.4%）、イタリア（1.4%）——は"
    "緩やかな人口転換を経験し大規模な移民ショックがない国々。",
    size=12, space_after=12,
)

# --- TABLE 1 ---
add_table_caption(
    doc,
    "表1. 40カ国のモデル適合度の要約"
    "（モデルバリアント・基準年別）",
)

table = doc.add_table(rows=6, cols=6)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = [
    'モデル', '期間(年)', 'N',
    'MAPE平均(%)', 'MAPE中央値(%)',
    '最終比率(mean ± SD)',
]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    set_cell_font(cell, 10)

data_rows = [
    ['静的 (1970)', '50', '40', '12.4', '7.3',
     '1.272 ± 0.481'],
    ['静的 (1980)', '43', '40', '9.6', '7.7',
     '1.023 ± 0.288'],
    ['静的 (1990)', '33', '40', '7.8', '6.5',
     '0.953 ± 0.198'],
    ['静的 (2000)', '23', '40', '5.1', '4.7',
     '0.914 ± 0.101'],
    ['動的 (10年)', '53', '40', '6.7', '4.6',
     '0.999 ± 0.189'],
]
for i, row_data in enumerate(data_rows):
    for j, val in enumerate(row_data):
        cell = table.rows[i + 1].cells[j]
        cell.text = val
        set_cell_font(cell, 10)

add_para(doc, '', size=6, space_after=6)

# --- 3.2 ---
add_heading_styled(doc, '3.2 不適合の原因', level=2)

add_para(
    doc,
    "MAPEが10%を超える国には共通の特徴がある。"
    "移民主導の人口増加はオーストラリア（13.5%）、"
    "カナダ（12.2%）、スイス（7.2%）、"
    "ルクセンブルク（21.5%）、イスラエル（13.9%）の"
    "不適合を説明する。本モデルはカンタム・テンポ・生存"
    "メカニズムを分離するため意図的に移民を除外しており、"
    "残差的不適合は移民成分を定量化する。",
    size=12, space_after=6,
)

add_para(
    doc,
    "急速な出生率転換がメキシコ（23.3%）、"
    "トルコ（17.0%）、中国（15.6%）、"
    "コロンビア（13.1%）を説明。"
    "韓国（11.9%）は極端な出生率低下と近年の移民の両効果を"
    "併せ持つ。リトアニア（12.0%）とラトビア（8.5%）は"
    "EU加盟後の移民流出を反映。",
    size=12, space_after=6,
)

add_para(
    doc,
    "図1は代表6カ国のモデル適合を示す。"
    "図2は全40カ国の検証結果、図3は基準年別の"
    "静的モデル性能、図4は静的・動的モデルの比較、"
    "図5はモデルバイアス分析を示す。",
    size=12, space_after=12,
)

# --- FIGURES 1-5 (inline after first-mention paragraph) ---
add_figure(
    doc, os.path.join(FIG_DIR, 'fig1_showcase.png'),
    "図1. 代表6カ国のモデル vs 観測人口軌跡"
    "（1970–2023）。動的モデル（青破線）は"
    "10年ごとにパラメータ更新。静的モデル"
    "（赤点線）は1970年基準。黒実線＝UN WPP 2024。",
    width=6.0,
)

add_figure(
    doc, os.path.join(FIG_DIR, 'fig2_all_countries.png'),
    "図2. 全40カ国のモデル検証。各パネル"
    "右上に動的モデルのMAPEを表示。"
    "アルファベット順。",
    width=6.5,
)

add_figure(
    doc, os.path.join(FIG_DIR, 'fig3_heatmap.png'),
    "図3. 静的モデルMAPE（%）：国×基準年。"
    "緑＝良好、赤＝不良。スケール上限30%。",
    width=5.0,
)

add_figure(
    doc, os.path.join(FIG_DIR, 'fig4_comparison.png'),
    "図4. 静的 vs 動的モデル比較。左：国別MAPE。"
    "右：最終人口比率（2023年）。",
    width=6.0,
)

add_figure(
    doc, os.path.join(FIG_DIR, 'fig5_bias.png'),
    "図5. モデルバイアス分析（基準年2000）。"
    "(A) vs TFR、(B) vs 平均寿命、(C) vs MAC。"
    "体系的関係は観察されず、モデル性能が"
    "人口学的文脈に対し頑健であることを示す。",
    width=6.0,
)

# ===================================================================
# SECTION 4
# ===================================================================
doc.add_page_break()
add_heading_styled(
    doc,
    '4. 政策レバーとしてのテンポ効果：'
    '社会適応速度の制御',
    level=1,
)

add_para(
    doc,
    "検証結果は、4つのパラメータ——TFR、平均寿命、MAC、"
    "出生スケジュール幅——だけのモデルが観測された人口"
    "軌跡を誤差中央値5%未満で再現できることを確認する。"
    "この簡素さにより、カンタム（TFR）、生存（e₀）、"
    "テンポ（MAC）の人口規模へのそれぞれの寄与が透明になる。",
    size=12, space_after=12,
)

add_para(
    doc,
    "テンポの経路は世代の重なりを通じて作用する。"
    "平均出産年齢が25歳の場合、おおよそ4世代"
    "（0, 25, 50, 75歳）が同時に生存する。MACが30歳に"
    "上昇すると世代間隔が広がり約3.3の重なり世代となり、"
    "同時在生人口が約6分の1減少する——各女性の出生数が"
    "同じであっても。このメカニズムはGoldstein, Lutz, Scherbov"
    "（2003）がEU15カ国で初めて同定したものであり、"
    "本稿の40カ国分析がその一般性を確認する。",
    size=12, space_after=12,
)

add_para(
    doc,
    "このメカニズムには人口規模を超えて人口変動のペースに"
    "及ぶ決定的な政策的含意がある。TFR=1.5で同一だがMAC=25 vs "
    "MAC=33の2カ国を考える。MACが高い国は重なる世代が少ない"
    "ため暦年あたりの実効的な人口減少が速い。この加速は"
    "年金制度改革、医療インフラ拡充、労働市場再構築のために"
    "利用可能な時間を圧縮する。",
    size=12, space_after=12,
)

add_para(
    doc,
    "テンポに敏感な政策は単に何人が存在するかだけでなく、"
    "社会が人口変動に合わせてその社会構造を適応させなければ"
    "ならない速度を制御する。若い家庭向けの手頃な住宅、"
    "普遍的な保育、早期の出産を不利にしない教育課程の再構築"
    "を通じてAFBを緩やかに引き下げる政策は、TFRを引き上げなく"
    "とも人口減少のペースを緩め制度的調整のための時間を"
    "稼ぐことができる。",
    size=12, space_after=12,
)

add_para(
    doc,
    "問いは「どうすれば出生数を増やせるか」だけでなく"
    "「人口転換の速度をどう管理するか」でもある。"
    "テンポ効果は後者に対するメカニズムを提供し、"
    "それは現代の政策設計において見落とされてきた"
    "（Lutz, Skirbekk, Testa 2006; BongaartsとSobotka 2012）。",
    size=12, space_after=12,
)

add_para(
    doc,
    "40カ国の検証は、これが理論的な珍事ではなく、"
    "多様な人口学的文脈——転換後の日本"
    "（MAC=31.4、TFR=1.2）から転換前のDRC"
    "（MAC=24.8、TFR=6.1）まで——で作用する"
    "量的に有意な力であることを実証する。",
    size=12, space_after=12,
)

# ===================================================================
# SECTION 5
# ===================================================================
add_heading_styled(doc, '5. 限界', level=1)

add_para(
    doc,
    "いくつかの限界を認める。第一に、モデルは移民を除外して"
    "おり、これはオーストラリア、カナダ、ルクセンブルク等の"
    "主要な不適合原因。除外は自然な人口学的メカニズムを"
    "分離するための意図的なものだが、大規模な純移民のある"
    "国への直接的な適用を制限する。第二に、正規分布の"
    "出生スケジュールは簡略化であり実際のASFRは歪みや"
    "二峰性を示す場合がある。第三に、10年ごとのパラメータ"
    "更新は10年内の急速な転換を見落とす可能性がある。"
    "第四に、Gompertz生存関数は成人死亡率を良好に適合するが"
    "乳児・小児死亡率の近似精度はやや低い。これらの"
    "限界にもかかわらず、モデルの簡素な構造は政策"
    "コミュニケーションにおける利点である。",
    size=12, space_after=12,
)

# ===================================================================
# SECTION 6
# ===================================================================
add_heading_styled(doc, '6. 結論', level=1)

add_para(
    doc,
    "同時在生人口に対するテンポ効果は人口学理論において"
    "確立されている（BongaartsとFeeney 1998; Goldstein, Lutz, "
    "Scherbov 2003）が政策実践においては忘れ去られている。"
    "UN WPP 2024データ（United Nations 2024）に対し"
    "40カ国で検証した簡素なモデルを用いて、出産の"
    "タイミングが人口規模に量的に実質的な影響を及ぼす"
    "こと——それが出生率のカンタムとは独立に作用する"
    "こと——を示した。政策的含意は人口規模を超えて"
    "人口変動のペースに及ぶ：テンポに敏感な介入は"
    "社会が制度を適応させなければならない速度を制御でき、"
    "従来の少子化対策を補完するレバーを提供する。"
    "テンポ効果を人口学的影響評価と人口政策設計に"
    "体系的に組み込むことを提言する。",
    size=12, space_after=18,
)

# ===================================================================
# REFERENCES (alphabetical, author-date style per PDR guidelines)
# ===================================================================
add_heading_styled(doc, '参考文献', level=1)

pdr_refs = [
    'Bongaarts, John and Griffith Feeney. 1998. '
    '“On the quantum and tempo of fertility,” '
    'Population and Development Review 24(2): 271–291.',

    'Bongaarts, John and Tomáš Sobotka. 2012. '
    '“A demographic explanation for the recent rise in '
    'European fertility,” Population and Development Review '
    '38(1): 83–120.',

    'Goldstein, Joshua R., Wolfgang Lutz, and Sergei Scherbov. '
    '2003. “Long-term population decline in Europe: The '
    'relative importance of tempo effects and generational '
    'length,” Population and Development Review 29(4): '
    '699–707.',

    'Gonand, Frédéric. 2005. “Assessing the robustness '
    'of demographic projections in OECD countries,” OECD '
    'Economics Department Working Papers No. 464. Paris: OECD '
    'Publishing.',

    'Lutz, Wolfgang, Vegard Skirbekk, and Maria Rita Testa. '
    '2006. “The low-fertility trap hypothesis: Forces that '
    'may lead to further postponement and fewer births in '
    'Europe,” Vienna Yearbook of Population Research 4: '
    '167–192.',

    'Stevens, Gretchen A., Leontine Alkema, Robert E. Black, '
    'et al. 2016. “Guidelines for Accurate and Transparent '
    'Health Estimates Reporting: The GATHER statement,” '
    'The Lancet 388(10062): e19–e23.',

    'United Nations, Department of Economic and Social Affairs, '
    'Population Division. 2024. World Population Prospects '
    '2024. New York: United Nations. '
    'https://population.un.org/wpp/',
]

for r in pdr_refs:
    add_para(doc, r, size=11, space_after=4)

# ===================================================================
# APPENDIX A: GATHER Compliance Statement
# ===================================================================
doc.add_page_break()
add_heading_styled(
    doc, '付録A：GATHER準拠声明', level=2,
)

add_para(
    doc,
    "本研究は人口推計を報告するものであり、"
    "GATHER（Stevens et al. 2016）に準拠する。主要項目：",
    size=11, space_after=6,
)

gather_items = [
    "項目1–3（目的、方法、対象集団）："
    "第1–2節に記述。",

    "項目4–7（データ入力）：全入力データは"
    "UN WPP 2024（United Nations 2024）から取得、"
    "population.un.org/wppで公開入手可能。一次データ収集なし。",

    "項目8–10（データ調整）：初期人口年齢構造を"
    "5歳階級から1歳刻みに均一分布で内挿。",

    "項目11–13（モデリング）：Gompertz生存、"
    "正規出生スケジュール、内生更新を第2.1節に記述。"
    "期間あたり4パラメータ（TFR、e₀、MAC、σ）。",

    "項目14–16（不確実性、結果）：MAPEと最終比率を"
    "適合指標として報告。モデルは決定論的。",

    "項目17–18（解釈、再現性）：コードとデータ"
    "ソースを文書化。分析コードは要請に応じ提供。",
]
for item in gather_items:
    add_para(doc, "• " + item, size=11, space_after=4)

# ===================================================================
# APPENDIX B: National Projection Methods Comparison
# ===================================================================
doc.add_page_break()
add_heading_styled(
    doc,
    '\u4ed8\u9332B\uff1aOECD\u8af8\u56fd\u306e\u516c\u5f0f\u4eba\u53e3\u4e88\u6e2c\u624b\u6cd5\u3068'
    '\u4eee\u5b9a\u306e\u6bd4\u8f03',
    level=2,
)

add_para(
    doc,
    "\u672c\u4ed8\u9332\u306f\u3001\u5206\u6790\u5bfe\u8c6140\u30ab\u56fd\u306e\u516c\u5f0f\u4eba\u53e3\u4e88\u6e2c\u624b\u6cd5\u3068"
    "\u4e3b\u8981\u4eee\u5b9a\u3092\u8981\u7d04\u3059\u308b\u3002\u5168\u3066\u306e\u56fd\u304c\u30b3\u30fc\u30db\u30fc\u30c8\u8981\u56e0\u6cd5"
    "\uff08cohort-component method\uff09\u306e\u5909\u7a2e\u3092\u57fa\u790e\u3068\u3057\u3066\u3044\u308b\u304c\u3001"
    "\u51fa\u751f\u30bf\u30a4\u30df\u30f3\u30b0\u306e\u6271\u3044\u3001\u6b7b\u4ea1\u7387\u6539\u5584\u30e2\u30c7\u30eb\u3001"
    "\u79fb\u6c11\u4eee\u5b9a\u3001\u30b7\u30ca\u30ea\u30aa\u69cb\u9020\u306b\u304a\u3044\u3066\u5927\u304d\u304f\u7570\u306a\u308b"
    "\uff08Gonand 2005\uff09\u3002\u3053\u308c\u3089\u306e\u5dee\u7570\u306f\u3001\u672c\u30e2\u30c7\u30eb\u306e\u610f\u56f3\u7684\u306a"
    "4\u30d1\u30e9\u30e1\u30fc\u30bf\u3078\u306e\u7c21\u7d20\u5316\u3092\u6587\u8108\u5316\u3059\u308b\u3002",
    size=11, space_after=12,
)

# --- Table B1 ---
add_table_caption(
    doc,
    "\u8868B1. \u56fd\u30fb\u6a5f\u95a2\u5225\u306e\u516c\u5f0f\u4eba\u53e3\u4e88\u6e2c\u624b\u6cd5\u306e\u6982\u8981",
)

tbl = doc.add_table(rows=16, cols=5)
tbl.style = 'Light Shading Accent 1'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = [
    '\u56fd\u30fb\u6a5f\u95a2', '\u624b\u6cd5', '\u51fa\u751f\u7387\u4eee\u5b9a',
    '\u6b7b\u4ea1\u7387\u4eee\u5b9a', '\u79fb\u6c11\u306e\u6271\u3044',
]
for i, h in enumerate(hdr):
    cell = tbl.rows[0].cells[i]
    cell.text = h
    set_cell_font(cell, 9)

rows_data = [
    ['\u56fd\u9023WPP 2024\n\uff08\u5168\u5bfe\u8c61\u56fd\uff09',
     '\u30b3\u30fc\u30db\u30fc\u30c8\u8981\u56e0\u6cd5\n\u30d9\u30a4\u30ba\u78ba\u7387\u7684\u4e88\u6e2c',
     '\u30d9\u30a4\u30ba\u968e\u5c64\u30e2\u30c7\u30eb\nTFR\u8ecc\u8de1\uff0b\u4e0d\u78ba\u5b9f\u6027',
     'Lee-Carter\u5909\u7a2e\n\u56fd\u5225\u30c9\u30ea\u30d5\u30c8',
     '\u7d14\u79fb\u6c11\u3092\u4eee\u5b9a\n\u9577\u671f\u5e73\u5747\u306b\u53ce\u675f'],
    ['\u65e5\u672c\uff08\u793e\u4eba\u7814\uff09',
     '\u30b3\u30fc\u30db\u30fc\u30c8\u8981\u56e0\u6cd5\n\u51fa\u751f3\u00d7\u6b7b\u4ea13\u30d0\u30ea\u30a2\u30f3\u30c8',
     '\u30b3\u30fc\u30db\u30fc\u30c8\u51fa\u751f\u30e2\u30c7\u30eb\n\u4e2d\u4f4dTFR=1.20(2070)\nMAC=32.8',
     'Lee-Carter\u30e2\u30c7\u30eb\ne\u2080=85.9(M)/91.8(F)\n2070\u5e74',
     '\u5e74\u9f62\u30fb\u6027\u5225\u7d14\u79fb\u6c11\n\u7d0416.3\u4e07\u4eba/\u5e74'],
    ['\u7c73\u56fd\uff08Census Bureau\uff09',
     '\u30b3\u30fc\u30db\u30fc\u30c8\u8981\u56e0\u6cd5\n\u4e3b\uff0b\u79fb\u6c113\u30d0\u30ea\u30a2\u30f3\u30c8',
     '\u4eba\u7a2e\u5225ASFR\nTFR\u21921.75(2060)',
     '\u6b7b\u56e0\u5225\u30e2\u30c7\u30eb\ne\u2080\u224883.9(2100)',
     '4\u30b7\u30ca\u30ea\u30aa\n\u7d04110\u4e07\u4eba/\u5e74\uff08\u4e3b\uff09'],
    ['\u30c9\u30a4\u30c4\uff08Destatis\uff09',
     '\u30b3\u30fc\u30db\u30fc\u30c8\u8981\u56e0\u6cd5\n27\u30d0\u30ea\u30a2\u30f3\u30c8(3\u00d73\u00d73)',
     'TFR 1.29\u20131.65',
     'e\u2080 82.6\u201386.4(M)\n85.9\u201389.3(F)',
     '\u7d14\u79fb\u6c113\u6c34\u6e96\n15\u4e07/25\u4e07/35\u4e07\u4eba'],
    ['\u82f1\u56fd\uff08ONS\uff09',
     '\u30b3\u30fc\u30db\u30fc\u30c8\u8981\u56e0\u6cd5\n\u4e3b\uff0b9\u30d0\u30ea\u30a2\u30f3\u30c8',
     '\u4e3bTFR\u22481.59',
     '\u6b7b\u4ea1\u7387\u6539\u5584\u30e2\u30c7\u30eb\ne\u2080\u224883.9(M)/86.3(F)',
     '\u9577\u671f\u7d14\u79fb\u6c11\u224831.5\u4e07\n\u30d0\u30ea\u30a2\u30f3\u30c8:12.6-51.5\u4e07'],
    ['\u30d5\u30e9\u30f3\u30b9\uff08INSEE\uff09',
     '\u30b3\u30fc\u30db\u30fc\u30c8\u8981\u56e0\u6cd5\n\u4e2d\u592e\uff0b\u6210\u5206\u5225\u30d0\u30ea\u30a2\u30f3\u30c8',
     'TFR\u22481.80\u4e2d\u592e\n\u9ad82.10/\u4f4e1.60',
     '\u30c8\u30ec\u30f3\u30c9\u5916\u633f\ne\u2080\u224887.5(M)/90.0(F)',
     '\u7d14\u79fb\u6c11+7\u4e07\u4eba/\u5e74'],
    ['\u97d3\u56fd\uff08KOSTAT\uff09',
     '\u30b3\u30fc\u30db\u30fc\u30c8\u8981\u56e0\u6cd5\n3\u30b7\u30ca\u30ea\u30aa',
     '\u30b3\u30fc\u30db\u30fc\u30c8\u30e2\u30c7\u30eb\n\u4e2d\u4f4dTFR\u21921.08(2040)',
     'Lee-Carter\ne\u2080=88.0(M)/91.4(F)',
     '\u56fd\u7c4d\u5225\u7d14\u79fb\u6c11\n\u7d046-10\u4e07\u4eba/\u5e74'],
    ['\u30a4\u30bf\u30ea\u30a2\uff08ISTAT\uff09',
     '\u30b3\u30fc\u30db\u30fc\u30c8\u8981\u56e0\u6cd5\n\u4e2d\u4f4d\uff0b4\u30b7\u30ca\u30ea\u30aa',
     'TFR\u22481.40\u4e2d\u4f4d\n\u7bc4\u56f21.20\u20131.60',
     'Lee-Carter\ne\u2080\u224885.8(M)/89.2(F)',
     '\u7d14\u79fb\u6c11\u2248+15-23\u4e07\u4eba/\u5e74'],
    ['\u8c6a\u5dde\uff08ABS\uff09',
     '\u30b3\u30fc\u30db\u30fc\u30c8\u8981\u56e0\u6cd5\n3\u7cfb\u5217(A/B/C)',
     'TFR 1.55\u20131.85\n\u7cfb\u5217B: 1.62',
     '\u6b7b\u4ea1\u7387\u6539\u5584\u7387\u5916\u633f\ne\u2080\u224887(M)/89(F)',
     'NOM\u4f9d\u5b58\u5ea6\u9ad8\n\u7cfb\u5217B:\u7d0423.5\u4e07/\u5e74'],
    ['\u30ab\u30ca\u30c0\uff08StatCan\uff09',
     '\u30b3\u30fc\u30db\u30fc\u30c8\u8981\u56e0\u6cd5\n\uff0b\u30de\u30a4\u30af\u30ed\u30b7\u30df\u30e5\u30ec\u30fc\u30b7\u30e7\u30f3',
     'TFR 1.40\u20131.60\n\u4e2d\u4f4d1.49',
     'Lee-Carter\u5909\u7a2e\ne\u2080\u224886(M)/89(F)',
     '\u7d14\u79fb\u6c11\u7d0440-50\u4e07/\u5e74\n\u4e3b\u8981\u6210\u9577\u30c9\u30e9\u30a4\u30d0\u30fc'],
    ['Eurostat\n\uff08EU\u52a0\u76df\u56fd\uff09',
     '\u30b3\u30fc\u30db\u30fc\u30c8\u8981\u56e0\u6cd5\n\u53ce\u675f\u30e2\u30c7\u30eb',
     '\u52a0\u76df\u56fd\u9593TFR\n\u90e8\u5206\u53ce\u675f',
     '\u6b7b\u4ea1\u7387\u6539\u5584\u7387\n\u53ce\u675f\u30e2\u30c7\u30eb',
     '\u9577\u671f\u7d14\u79fb\u6c11\u306b\n\u53ce\u675f\u3059\u308b\u56fd\u5225\u7d4c\u8def'],
    ['\u4e2d\u56fd\uff08NBS\uff09',
     '\u30b3\u30fc\u30db\u30fc\u30c8\u8981\u56e0\u6cd5\n\uff08\u5b9a\u671f\u516c\u8868\u306a\u3057\uff09',
     'TFR=1.0-1.2(2022-23)\nUN WPP\u306f\u56de\u5fa9\u4eee\u5b9a',
     '\u30e2\u30c7\u30eb\u751f\u547d\u8868\ne\u2080\u224878.6',
     '\u4f4e\u3044\u56fd\u969b\u79fb\u6c11'],
    ['DRC\uff08\u56fd\u5bb6\u4e88\u6e2c\u306a\u3057\uff09',
     'UN WPP\u306b\u4f9d\u5b58\n\u72ec\u7acb\u4e88\u6e2c\u306a\u3057',
     'TFR\u22486.1(2023)\nUN:\u6f38\u6e1b\u4eee\u5b9a',
     '\u30e2\u30c7\u30eb\u751f\u547d\u8868\ne\u2080\u224860.7',
     '\u4f4e\u3044\u7d14\u79fb\u6c11\n\u96e3\u6c11\u6d41\u306f\u975e\u4f53\u7cfb\u7684'],
    ['\u30e1\u30ad\u30b7\u30b3\uff08CONAPO\uff09',
     '\u30b3\u30fc\u30db\u30fc\u30c8\u8981\u56e0\u6cd5\n3\u30d0\u30ea\u30a2\u30f3\u30c8',
     'TFR\u2192\u7d041.7(2050)',
     '\u30c8\u30ec\u30f3\u30c9\u5916\u633f\ne\u2080\u224879(M)/83(F)',
     '\u7d14\u79fb\u51fa\u2192\u307b\u307c\u30bc\u30ed\n\u22485\u4e07\u4eba(2050)'],
    ['\u30c8\u30eb\u30b3\uff08TurkStat\uff09',
     '\u30b3\u30fc\u30db\u30fc\u30c8\u8981\u56e0\u6cd5\n3\u30b7\u30ca\u30ea\u30aa',
     'TFR 1.51\u2192\u7d041.60\n\u9577\u671f',
     '\u6539\u5584\u30e2\u30c7\u30eb\ne\u2080\u224880(M)/84(F)',
     '\u7d14\u79fb\u5165\u224820-30\u4e07/\u5e74\n\u96e3\u6c11\u542b\u3080'],
]
for i, rd in enumerate(rows_data):
    for j, val in enumerate(rd):
        cell = tbl.rows[i + 1].cells[j]
        cell.text = val
        set_cell_font(cell, 9)

add_para(doc, '', size=6, space_after=8)

add_para(
    doc,
    "\u51fa\u5178\uff1aUnited Nations (2024), \u793e\u4eba\u7814 (2023), US Census "
    "Bureau (2023), Destatis (2025), ONS (2025), INSEE "
    "(2021), KOSTAT (2023), ISTAT (2023), ABS (2018), "
    "Statistics Canada (2024), Eurostat (2024), "
    "CONAPO (2018), TurkStat (2023)\u3002",
    italic=True, size=9, space_after=12,
)

# --- B.1 ---
add_heading_styled(
    doc, 'B.1 \u5171\u901a\u70b9\u3068\u4e3b\u8981\u306a\u5dee\u7570', level=3,
)

add_para(
    doc,
    "\u5168\u3066\u306e\u56fd\u5bb6\u4e88\u6e2c\u30b7\u30b9\u30c6\u30e0\u304c\u30b3\u30fc\u30db\u30fc\u30c8\u8981\u56e0\u6cd5\u3092"
    "\u57fa\u790e\u69cb\u9020\u3068\u3057\u3066\u5171\u6709\u3057\u3001\u5e74\u9f62\u5225\u4eba\u53e3\u3092\u51fa\u751f\u30fb\u6b7b\u4ea1\u30fb"
    "\u79fb\u6c11\u306e\u4eee\u5b9a\u3092\u7528\u3044\u3066\u53cd\u5fa9\u7684\u306b\u52a0\u9f62\u3055\u305b\u308b\u3002"
    "\u672c\u30e2\u30c7\u30eb\u3068\u306e\u95a2\u9023\u3067\u91cd\u8981\u306a\u5dee\u7570\u306f\u4ee5\u4e0b\u306e\u901a\u308a\uff1a",
    size=11, space_after=6,
)

bullet_texts = [
    "\u51fa\u751f\u30bf\u30a4\u30df\u30f3\u30b0\u306e\u6271\u3044\uff1a\u591a\u304f\u306e\u56fd\u5bb6\u4e88\u6e2c\u306fMAC\u3068\u03c3\u3067"
    "\u30d1\u30e9\u30e1\u30fc\u30bf\u5316\u3059\u308b\u306e\u3067\u306f\u306a\u304f\u5b8c\u5168\u306aASFR\u30b9\u30b1\u30b8\u30e5\u30fc\u30eb\u3092"
    "\u6307\u5b9a\u3059\u308b\u3002\u65e5\u672c\uff08\u793e\u4eba\u7814\uff09\u3068\u97d3\u56fd\u306f\u30bf\u30a4\u30df\u30f3\u30b0\u30b7\u30d5\u30c8\u3092"
    "\u660e\u793a\u7684\u306b\u8ffd\u8de1\u3059\u308b\u30b3\u30fc\u30db\u30fc\u30c8\u51fa\u751f\u30e2\u30c7\u30eb\u3092\u4f7f\u7528\u3002"
    "\u672c\u30e2\u30c7\u30eb\u306e\u6b63\u898f\u5206\u5e03\u7c21\u7565\u5316\u306f\u4e2d\u5fc3\u50be\u5411\u3092\u6355\u6349\u3059\u308b\u304c"
    "\u30b9\u30b1\u30b8\u30e5\u30fc\u30eb\u5f62\u72b6\u306f\u6355\u3048\u306a\u3044\u3002",

    "\u6b7b\u4ea1\u7387\u30e2\u30c7\u30eb\uff1a\u56fd\u5bb6\u6a5f\u95a2\u306f\u901a\u5e38Lee-Carter\u307e\u305f\u306f\u305d\u306e"
    "\u62e1\u5f35\u3092\u4f7f\u7528\u3002\u672c\u30e2\u30c7\u30eb\u306eGompertz\u751f\u5b58\uff08\u30d1\u30e9\u30e1\u30fc\u30bfa\u306e\u307f"
    "\u30ad\u30e3\u30ea\u30d6\u30ec\u30fc\u30b7\u30e7\u30f3\uff09\u306f\u3088\u308a\u7c21\u7d20\u3060\u304c\u5e74\u9f62\u5225\u6b7b\u4ea1\u7387"
    "\u30d1\u30bf\u30fc\u30f3\u306e\u67d4\u8edf\u6027\u306f\u4f4e\u3044\u3002",

    "\u79fb\u6c11\uff1a\u6700\u3082\u5909\u52d5\u304c\u5927\u304d\u3044\u6210\u5206\u3067\u3042\u308a\u672c\u30e2\u30c7\u30eb\u304c"
    "\u610f\u56f3\u7684\u306b\u9664\u5916\u3059\u308b\u3082\u306e\u3002\u79fb\u6c11\u4f9d\u5b58\u56fd\uff08\u8c6a\u5dde\u3001\u30ab\u30ca\u30c0\u3001"
    "\u30eb\u30af\u30bb\u30f3\u30d6\u30eb\u30af\u3001\u30a4\u30b9\u30e9\u30a8\u30eb\uff09\u3067\u306f\u79fb\u6c11\u4eee\u5b9a\u304c"
    "\u9577\u671f\u4e88\u6e2c\u3092\u652f\u914d\u3059\u308b\u3002",

    "\u30b7\u30ca\u30ea\u30aa\u69cb\u9020\uff1a\u97d3\u56fd\u306e3\u30d0\u30ea\u30a2\u30f3\u30c8\u304b\u3089\u30c9\u30a4\u30c4\u306e"
    "27\u30d0\u30ea\u30a2\u30f3\u30c8\u307e\u3067\u5e45\u5e83\u3044\u3002UN WPP\u306f\u30d9\u30a4\u30ba\u78ba\u7387\u7684\u4e88\u6e2c\u3067"
    "\u5b8c\u5168\u306a\u4e0d\u78ba\u5b9f\u6027\u5206\u5e03\u3092\u63d0\u4f9b\u3002\u672c\u30e2\u30c7\u30eb\u306e\u6c7a\u5b9a\u8ad6\u7684"
    "\u5358\u4e00\u8ecc\u8de1\u306f\u4e0d\u78ba\u5b9f\u6027\u5b9a\u91cf\u5316\u3092\u30c6\u30f3\u30dd\u30fb\u30ab\u30f3\u30bf\u30e0\u5206\u89e3\u306e"
    "\u900f\u660e\u6027\u3068\u5f15\u304d\u63db\u3048\u306b\u3057\u3066\u3044\u308b\u3002",

    "\u30c6\u30f3\u30dd\u306e\u6271\u3044\uff1a\u6ce8\u76ee\u3059\u3079\u304d\u3053\u3068\u306b\u3001\u3044\u305a\u308c\u306e\u56fd\u5bb6"
    "\u4e88\u6e2c\u30b7\u30b9\u30c6\u30e0\u3082\u30ab\u30f3\u30bf\u30e0\u3068\u30c6\u30f3\u30dd\u6210\u5206\u3078\u306e\u660e\u793a\u7684"
    "\u5206\u89e3\u3092\u884c\u3063\u3066\u3044\u306a\u3044\u3002\u51fa\u751f\u30bf\u30a4\u30df\u30f3\u30b0\u306fASFR\u3092\u901a\u3058\u3066"
    "\u6697\u9ed9\u7684\u306b\u5165\u308b\u304c\u3001MAC\u306e\u540c\u6642\u5728\u751f\u4eba\u53e3\u3078\u306e\u72ec\u7acb\u3057\u305f"
    "\u5bc4\u4e0e\u306f\u5206\u96e2\u3055\u308c\u306a\u3044\u3002\u3053\u306e\u30ae\u30e3\u30c3\u30d7\u304c\u672c\u7814\u7a76\u306e\u52d5\u6a5f\u3067\u3042\u308b\u3002",
]
for bt in bullet_texts:
    add_para(doc, "\u2022 " + bt, size=11, space_after=4)

# --- B.2 ---
add_heading_styled(
    doc, 'B.2 \u30e2\u30c7\u30eb\u6bd4\u8f03\u3078\u306e\u542b\u610f', level=3,
)

add_para(
    doc,
    "\u672c\u30e2\u30c7\u30eb\u306f\u56fd\u5bb6\u4e88\u6e2c\u30b7\u30b9\u30c6\u30e0\u306e\u4ee3\u66ff\u3067\u306f\u306a\u304f\u3001"
    "\u30c6\u30f3\u30dd\u30fb\u30ab\u30f3\u30bf\u30e0\u30fb\u751f\u5b58\u5206\u89e3\u3092\u660e\u793a\u3059\u308b\u3053\u3068\u306b\u3088\u308b"
    "\u88dc\u5b8c\u3092\u76ee\u7684\u3068\u3059\u308b\u3002\u4e0a\u8868\u306f\u3001\u6700\u3082\u7cbe\u7dfb\u306a\u56fd\u5bb6\u30b7\u30b9\u30c6\u30e0\u3067"
    "\u3055\u3048\u540c\u3058\u57fa\u672c\u69cb\u9020\uff08\u30b3\u30fc\u30db\u30fc\u30c8\u8981\u56e0\u6cd5\uff09\u3092\u5171\u6709\u3057\u3001"
    "\u30d1\u30e9\u30e1\u30fc\u30bf\u63a8\u5b9a\u6cd5\u3068\u30b7\u30ca\u30ea\u30aa\u69cb\u9020\u306b\u304a\u3044\u3066\u4e3b\u306b\u7570\u306a\u308a\u3001"
    "\u660e\u793a\u7684\u30c6\u30f3\u30dd\u5206\u89e3\u3092\u4e00\u69d8\u306b\u6b20\u304f\u3053\u3068\u3092\u793a\u3059\u3002"
    "4\u30d1\u30e9\u30e1\u30fc\u30bf\u30e2\u30c7\u30eb\u304c\u540c\u3058\u4eba\u53e3\u306b\u5bfe\u3057\u52d5\u7684MAPE\u4e2d\u592e\u5024"
    "4.6%\u3092\u9054\u6210\u3059\u308b\u6027\u80fd\u306f\u3001\u79fb\u6c11\u3092\u542b\u3080\u5b8c\u5168\u30d1\u30e9\u30e1\u30fc\u30bf\u5316"
    "\u56fd\u5bb6\u30e2\u30c7\u30eb\u306e\u7cbe\u5ea6\u306b\u306f\u53ca\u3070\u306a\u3044\u3082\u306e\u306e\u3001\u30c6\u30f3\u30dd\u7d4c\u8def\u306e"
    "\u91cf\u7684\u91cd\u8981\u6027\u3092\u78ba\u7acb\u3059\u308b\u306b\u306f\u5341\u5206\u3067\u3042\u308b\u3002",
    size=11, space_after=12,
)

# ===================================================================
# SAVE
# ===================================================================
outpath = os.path.join(OUT_DIR, 'PDR_Research_Note_JP.docx')
doc.save(outpath)
print(f'OK: {outpath}')
